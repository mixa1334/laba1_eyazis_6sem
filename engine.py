from nltk.tokenize import word_tokenize
from pymystem3 import Mystem
from string import punctuation
import docx
import pickle

import Vocabulary


def read_text_from_file(filename):
    doc = docx.Document(filename)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)


def write_text_to_file(filename, voc):
    doc_file = docx.Document()
    doc_file.add_heading("Словарь", 1)
    for lemma in voc.get_all_lemmas():
        doc_file.add_paragraph("---------------------------------------------------------------------------------")
        p = doc_file.add_paragraph()
        runner = p.add_run("Лексема -> " + str(lemma))
        runner.bold = True
        for word in voc.get_all_words_according_to_lemma(lemma):
            info = voc.get_morphological_info_according_to_word(word)
            word_text = "\tсловоформа -> " + str(word) + ", кол-во вхождения в текст -> " + str(
                voc.get_count_of_word(word, lemma)) + " доп информация -> часть речи: " + str(
                info[0]) + ", род: " + str(info[1]) + ", число: " + str(info[2]) + ", пажед: " + str(info[3])
            doc_file.add_paragraph(word_text)
    doc_file.save(filename + ".doc")


def write_vocabulary_to_file(voc, filename):
    with open(filename + ".pkl", "wb") as file:
        pickle.dump(voc, file)


def read_vocabulary_from_file(filename):
    with open(filename, "rb") as file:
        voc = pickle.load(file)
    return list(voc)


def preprocess_text(text):
    mystem = Mystem()
    text = text.lower()
    result = Vocabulary.Vocabulary()
    words = list(filter(lambda val: val not in punctuation, word_tokenize(text)))
    lemmas = list(filter(lambda val: val != " " and val != "\n", mystem.lemmatize(" ".join(words))))

    for i in range(len(words)):
        result.add_word_according_to_lemma(words[i], lemmas[i])

    return result


def process_word_to_lemma(word):
    mystem = Mystem()
    word = word.lower()
    lemma = mystem.lemmatize(word)
    return lemma
